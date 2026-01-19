import io
import os
import re
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


_HEADING_RE = re.compile(r"^(#{1,6})\\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\\d+\\.\\s+(.*)$")
_TABLE_SEPARATOR_RE = re.compile(r"^:?-{3,}:?$")


def _split_table_row(line: str) -> list[str]:
    trimmed = line.strip()
    if trimmed.startswith("|"):
        trimmed = trimmed[1:]
    if trimmed.endswith("|"):
        trimmed = trimmed[:-1]
    return [cell.strip() for cell in trimmed.split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    if not cells:
        return False
    return all(_TABLE_SEPARATOR_RE.match(cell) for cell in cells)


def _get_table_alignments(line: str) -> list[str]:
    alignments = []
    for cell in _split_table_row(line):
        left = cell.startswith(":")
        right = cell.endswith(":")
        if left and right:
            alignments.append("center")
        elif right:
            alignments.append("right")
        elif left:
            alignments.append("left")
        else:
            alignments.append("left")
    return alignments


def _normalize_markdown(markdown_text: str) -> str:
    text = markdown_text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?<!^)(?<!\n)(#{1,6}\s+)", r"\n\\1", text)
    text = re.sub(r"(?<!^)(?<!\n)\s+([-*+]\s+)", r"\n\\1", text)
    text = re.sub(r"(?<!^)(?<!\n)\s+(\d+\.\s+)", r"\n\\1", text)
    text = re.sub(r"\|\s*\|", "|\n|", text)
    return text.strip()


def _set_cell_alignment(cell, alignment: str) -> None:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para_alignment = mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in cell.paragraphs:
        paragraph.alignment = para_alignment


def _bold_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _pandoc_convert_to_docx(content: str, input_format: str) -> bytes | None:
    try:
        import pypandoc

        pypandoc.get_pandoc_version()
    except (ImportError, OSError):
        return None

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_file.close()
    try:
        pypandoc.convert_text(
            content,
            "docx",
            format=input_format,
            outputfile=tmp_file.name,
        )
        with open(tmp_file.name, "rb") as handle:
            return handle.read()
    except Exception:
        return None
    finally:
        try:
            os.unlink(tmp_file.name)
        except OSError:
            pass


def html_to_docx_bytes(html: str, fallback_markdown: str | None = None) -> bytes:
    pandoc_bytes = _pandoc_convert_to_docx(html, "html")
    if pandoc_bytes is not None:
        return pandoc_bytes
    if fallback_markdown is not None:
        return markdown_to_docx_bytes(fallback_markdown)
    return markdown_to_docx_bytes(html)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    pandoc_bytes = _pandoc_convert_to_docx(markdown_text, "gfm")
    if pandoc_bytes is not None:
        return pandoc_bytes

    doc = Document()
    paragraph_lines: list[str] = []
    lines = _normalize_markdown(markdown_text).splitlines()
    idx = 0

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            doc.add_paragraph(" ".join(paragraph_lines).strip())
            paragraph_lines = []

    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            flush_paragraph()
            idx += 1
            continue

        if (
            "|" in line
            and idx + 1 < len(lines)
            and _is_table_separator(lines[idx + 1].strip())
        ):
            flush_paragraph()
            header_cells = _split_table_row(line)
            alignments = _get_table_alignments(lines[idx + 1].strip())
            rows = []
            idx += 2

            while idx < len(lines):
                row_line = lines[idx].strip()
                if not row_line or "|" not in row_line:
                    break
                rows.append(_split_table_row(row_line))
                idx += 1

            col_count = max(len(header_cells), max((len(r) for r in rows), default=0))
            if col_count == 0:
                continue

            table = doc.add_table(rows=1 + len(rows), cols=col_count)
            table.style = "Table Grid"
            for col_idx in range(col_count):
                header_cell = table.cell(0, col_idx)
                header_cell.text = (
                    header_cells[col_idx] if col_idx < len(header_cells) else ""
                )
                _bold_cell(header_cell)
                _set_cell_alignment(
                    header_cell,
                    alignments[col_idx] if col_idx < len(alignments) else "left",
                )

            for row_idx, row in enumerate(rows, start=1):
                for col_idx in range(col_count):
                    data_cell = table.cell(row_idx, col_idx)
                    data_cell.text = (
                        row[col_idx] if col_idx < len(row) else ""
                    )
                    _set_cell_alignment(
                        data_cell,
                        alignments[col_idx] if col_idx < len(alignments) else "left",
                    )

            idx += 1
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 6)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            idx += 1
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            flush_paragraph()
            doc.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            idx += 1
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            flush_paragraph()
            doc.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            idx += 1
            continue

        paragraph_lines.append(line)
        idx += 1

    flush_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
